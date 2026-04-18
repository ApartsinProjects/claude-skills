#!/usr/bin/env python3
"""
Apply academic paper formatting to DOCX.

Applies:
- template-aligned manuscript styles
- centered front matter
- left-aligned figure/table captions
- journal-style tables
- reference hanging indents
- page numbers in the footer
- pagination controls for headings and captions

Usage: python apply_academic_style.py [--input input.docx] [--output output.docx]
"""

import argparse
import os
import re
import sys

try:
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt
except ImportError:
    print("Installing python-docx...")
    os.system("pip install python-docx")
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt


PROFILES = {
    "camera-ready-generic": {
        "font_family": "Times New Roman",
        "body_font_size": 11,
        "title_font_size": 18,
        "author_font_size": 12,
        "affiliation_font_size": 10,
        "caption_font_size": 9.5,
        "footer_font_size": 10,
        "reference_font_size": 10,
        "heading1_size": 14,
        "heading2_size": 12,
        "heading3_size": 11,
        "line_spacing": 1.15,
        "margin_inches": 1.0,
        "table_header_fill": "EEF3F7",
        "table_rule_dark": "4F6272",
        "table_rule_mid": "9FB0C2",
        "table_rule_light": "D9E1E8",
        "equation_space_before": 4,
        "equation_space_after": 4,
        "table_space_before_caption": 8,
        "table_space_after": 8,
    },
    "review-manuscript": {
        "font_family": "Times New Roman",
        "body_font_size": 11,
        "title_font_size": 18,
        "author_font_size": 12,
        "affiliation_font_size": 10,
        "caption_font_size": 10,
        "footer_font_size": 10,
        "reference_font_size": 10,
        "heading1_size": 14,
        "heading2_size": 12,
        "heading3_size": 11,
        "line_spacing": 1.5,
        "margin_inches": 1.0,
        "table_header_fill": "EEF3F7",
        "table_rule_dark": "4F6272",
        "table_rule_mid": "9FB0C2",
        "table_rule_light": "D9E1E8",
        "equation_space_before": 6,
        "equation_space_after": 6,
        "table_space_before_caption": 8,
        "table_space_after": 8,
    },
}


def ensure_style(styles, name, style_type, base=None):
    try:
        style = styles[name]
    except KeyError:
        style = styles.add_style(name, style_type)
    if base is not None:
        style.base_style = styles[base]
    return style


def clear_paragraph(paragraph):
    p = paragraph._element
    for child in list(p):
        p.remove(child)


def add_page_number_run(paragraph, config):
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"

    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run = paragraph.add_run()
    run.font.name = config["font_family"]
    run.font.size = Pt(config["footer_font_size"])
    r = run._r
    r.append(fld_begin)
    r.append(instr)
    r.append(fld_separate)

    display = OxmlElement("w:t")
    display.text = "1"
    r.append(display)
    r.append(fld_end)


def format_section_footer(section, config):
    footer = section.footer
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    clear_paragraph(paragraph)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    add_page_number_run(paragraph, config)


def set_spacing(paragraph, *, before=None, after=None, line_spacing=None):
    pf = paragraph.paragraph_format
    if before is not None:
        pf.space_before = Pt(before)
    if after is not None:
        pf.space_after = Pt(after)
    if line_spacing is not None:
        pf.line_spacing = line_spacing


def set_paragraph_flag(paragraph, tag_name, val="1"):
    ppr = paragraph._element.get_or_add_pPr()
    tag = ppr.find(qn(f"w:{tag_name}"))
    if tag is None:
        tag = OxmlElement(f"w:{tag_name}")
        ppr.append(tag)
    tag.set(qn("w:val"), str(val))


def apply_pagination_controls(paragraph, *, widow=True, keep_with_next=False, keep_together=False):
    set_paragraph_flag(paragraph, "widowControl", "1" if widow else "0")
    paragraph.paragraph_format.keep_with_next = keep_with_next
    paragraph.paragraph_format.keep_together = keep_together


def set_table_border(table, config):
    tbl = table._element
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)

    existing = tbl_pr.find(qn("w:tblBorders"))
    if existing is not None:
        tbl_pr.remove(existing)

    tbl_borders = OxmlElement("w:tblBorders")

    def border(name, val, size, color):
        elem = OxmlElement(f"w:{name}")
        elem.set(qn("w:val"), val)
        if val != "nil":
            elem.set(qn("w:sz"), str(size))
            elem.set(qn("w:color"), color)
        tbl_borders.append(elem)

    border("top", "single", 10, config["table_rule_dark"])
    border("bottom", "single", 10, config["table_rule_dark"])
    border("insideH", "single", 5, config["table_rule_mid"])
    border("left", "single", 6, config["table_rule_mid"])
    border("right", "single", 6, config["table_rule_mid"])
    border("insideV", "single", 5, config["table_rule_light"])

    tbl_pr.append(tbl_borders)


def set_table_width(table, width_percent=100):
    tbl = table._element
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "pct")
    tbl_w.set(qn("w:w"), str(width_percent * 50))


def set_table_layout(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl = table._element
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "autofit")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, config, bottom_color=None, bottom_size=6):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)

    for side in ("left", "right", "top"):
        edge = tc_borders.find(qn(f"w:{side}"))
        if edge is None:
            edge = OxmlElement(f"w:{side}")
            tc_borders.append(edge)
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), "4")
        edge.set(qn("w:color"), config["table_rule_light"])

    bottom = tc_borders.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        tc_borders.append(bottom)

    if bottom_color is None:
        bottom.set(qn("w:val"), "nil")
    else:
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), str(bottom_size))
        bottom.set(qn("w:color"), bottom_color)


def is_numericish(text):
    cleaned = text.strip().replace("%", "").replace(",", "")
    return bool(cleaned) and bool(re.fullmatch(r"[-+]?\d+(\.\d+)?", cleaned))


def detect_column_alignments(table):
    alignments = []
    for col_idx in range(len(table.columns)):
        values = []
        for row in table.rows[1:]:
            if col_idx < len(row.cells):
                txt = row.cells[col_idx].text.strip()
                if txt:
                    values.append(txt)
        if not values:
            alignments.append(WD_PARAGRAPH_ALIGNMENT.LEFT)
            continue
        short_ratio = sum(1 for v in values if len(v) <= 12) / len(values)
        numeric_ratio = sum(1 for v in values if is_numericish(v)) / len(values)
        if numeric_ratio >= 0.6 or (col_idx == 0 and short_ratio >= 0.8):
            alignments.append(WD_PARAGRAPH_ALIGNMENT.CENTER)
        else:
            alignments.append(WD_PARAGRAPH_ALIGNMENT.LEFT)
    return alignments


def format_table_cell(cell, config, *, is_header=False, alignment=WD_PARAGRAPH_ALIGNMENT.CENTER,
                      bottom_color=None, bottom_size=4):
    if is_header:
        set_cell_shading(cell, config["table_header_fill"])
    set_cell_border(cell, config, bottom_color=bottom_color, bottom_size=bottom_size)

    for para in cell.paragraphs:
        para.alignment = alignment
        set_spacing(para, before=0, after=1, line_spacing=1.0)
        apply_pagination_controls(para, widow=True, keep_with_next=False, keep_together=True)
        for run in para.runs:
            run.font.name = config["font_family"]
            run.font.size = Pt(10 if not is_header else 9.5)
            if is_header:
                run.font.bold = True


def has_drawing(element):
    for elem in element.iter():
        if "drawing" in elem.tag.lower():
            return True
    return False


def has_display_math(paragraph):
    xml = paragraph._element.xml
    return "<m:oMathPara" in xml or ("<m:oMath" in xml and not paragraph.text.strip())


def previous_block_is_table(paragraph):
    prev = paragraph._element.getprevious()
    return prev is not None and prev.tag == qn("w:tbl")


def emphasize_caption_label(paragraph):
    text = paragraph.text.strip()
    match = re.match(r"^((Figure|Table)\s+\d+:)(\s*)(.*)$", text)
    if not match:
        return

    label, _, spacer, rest = match.groups()
    clear_paragraph(paragraph)

    run = paragraph.add_run(label)
    run.bold = True
    run.italic = True

    if spacer:
        spacer_run = paragraph.add_run(spacer)
        spacer_run.bold = False
        spacer_run.italic = True

    if rest:
        rest_run = paragraph.add_run(rest)
        rest_run.bold = False
        rest_run.italic = True


def is_numbered_section_heading(text):
    return bool(re.match(r"^\s*\d+(\.\d+)*\.?\s+", text))


def is_caption_paragraph(text):
    return bool(re.match(r"^\s*(Figure|Table)\s+\d+[:.]", text))


def configure_base_styles(doc, config):
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = config["font_family"]
    normal.font.size = Pt(config["body_font_size"])
    normal.paragraph_format.line_spacing = config["line_spacing"]
    normal.paragraph_format.space_after = Pt(0)

    for style_name, size in (
        ("Title", config["title_font_size"]),
        ("Subtitle", config["author_font_size"]),
        ("Heading 1", config["heading1_size"]),
        ("Heading 2", config["heading2_size"]),
        ("Heading 3", config["heading3_size"]),
    ):
        style = styles[style_name]
        style.font.name = config["font_family"]
        style.font.size = Pt(size)

    for custom_name, base_name in (
        ("Abstract Label", "Body Text"),
        ("Keywords", "Body Text"),
        ("References Heading", "Heading 1"),
        ("Reference Entry", "Body Text"),
    ):
        ensure_style(styles, custom_name, WD_STYLE_TYPE.PARAGRAPH, base=base_name)

    for style_name in (
        "Body Text",
        "First Paragraph",
        "Compact",
        "Image Caption",
        "Table Caption",
        "Abstract Label",
        "Keywords",
        "Reference Entry",
    ):
        style = styles[style_name]
        style.font.name = config["font_family"]

    styles["Body Text"].font.size = Pt(config["body_font_size"])
    styles["Body Text"].paragraph_format.line_spacing = config["line_spacing"]
    styles["Body Text"].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    styles["First Paragraph"].font.size = Pt(config["body_font_size"])
    styles["First Paragraph"].paragraph_format.line_spacing = config["line_spacing"]
    styles["First Paragraph"].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    styles["Compact"].font.size = Pt(config["body_font_size"])
    styles["Compact"].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    styles["Compact"].paragraph_format.left_indent = Inches(0.22)
    styles["Compact"].paragraph_format.first_line_indent = Inches(-0.18)
    styles["Compact"].paragraph_format.line_spacing = 1.05
    styles["Compact"].paragraph_format.space_after = Pt(2)

    styles["Image Caption"].font.size = Pt(config["caption_font_size"])
    styles["Image Caption"].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    styles["Image Caption"].paragraph_format.line_spacing = 1.0
    styles["Image Caption"].paragraph_format.space_before = Pt(4)
    styles["Image Caption"].paragraph_format.space_after = Pt(6)

    styles["Table Caption"].font.size = Pt(config["caption_font_size"])
    styles["Table Caption"].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    styles["Table Caption"].paragraph_format.line_spacing = 1.0
    styles["Table Caption"].paragraph_format.space_before = Pt(config["table_space_before_caption"])
    styles["Table Caption"].paragraph_format.space_after = Pt(3)

    styles["Abstract Label"].font.size = Pt(config["body_font_size"])
    styles["Abstract Label"].font.bold = True
    styles["Abstract Label"].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    styles["Abstract Label"].paragraph_format.space_after = Pt(4)

    styles["Keywords"].font.size = Pt(config["body_font_size"])
    styles["Keywords"].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    styles["Keywords"].paragraph_format.space_after = Pt(4)

    styles["Reference Entry"].font.size = Pt(config["reference_font_size"])
    styles["Reference Entry"].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    styles["Reference Entry"].paragraph_format.left_indent = Inches(0)
    styles["Reference Entry"].paragraph_format.first_line_indent = Inches(0)
    styles["Reference Entry"].paragraph_format.line_spacing = 1.0
    styles["Reference Entry"].paragraph_format.space_after = Pt(2)

    styles["References Heading"].font.size = Pt(config["heading1_size"])
    styles["References Heading"].font.bold = True

    try:
        hyperlink_style = styles["Hyperlink"]
        hyperlink_style.font.color.rgb = None
        hyperlink_style.font.underline = False
        hyperlink_style.font.name = config["font_family"]
    except KeyError:
        pass


def classify_front_matter(doc):
    nonempty = [p for p in doc.paragraphs if p.text.strip()]
    manuscript_title = None
    if nonempty:
        for para in nonempty:
            if para.style and para.style.name.startswith("Heading") and not is_numbered_section_heading(para.text):
                manuscript_title = para
                break
        if manuscript_title is None:
            manuscript_title = nonempty[0]

    author_para = None
    affiliation_para = None
    if manuscript_title in nonempty:
        idx = nonempty.index(manuscript_title)
        for para in nonempty[idx + 1:]:
            text = para.text.strip()
            if text.lower().startswith("keywords:") or text.lower() == "abstract" or is_numbered_section_heading(text):
                break
            if author_para is None:
                author_para = para
                continue
            if affiliation_para is None:
                affiliation_para = para
                break

    return {
        "title": manuscript_title._element if manuscript_title is not None else None,
        "author": author_para._element if author_para is not None else None,
        "affiliation": affiliation_para._element if affiliation_para is not None else None,
    }


def apply_academic_formatting(doc, profile="camera-ready-generic", table_width=100):
    config = PROFILES[profile]

    print("Applying academic formatting...")
    configure_base_styles(doc, config)

    for section in doc.sections:
        section.top_margin = Inches(config["margin_inches"])
        section.bottom_margin = Inches(config["margin_inches"])
        section.left_margin = Inches(config["margin_inches"])
        section.right_margin = Inches(config["margin_inches"])
        format_section_footer(section, config)

    front = classify_front_matter(doc)

    image_count = 0
    text_count = 0
    in_references = False

    for para in doc.paragraphs:
        text = para.text.strip()
        style_name = para.style.name if para.style else ""
        has_image = has_drawing(para._element)
        is_heading = style_name.startswith("Heading")
        is_title = front["title"] is not None and para._element is front["title"]
        is_author = front["author"] is not None and para._element is front["author"]
        is_affiliation = front["affiliation"] is not None and para._element is front["affiliation"]
        is_caption = is_caption_paragraph(text)
        is_table_caption = text.startswith("Table ")
        is_figure_caption = text.startswith("Figure ")
        is_abstract_label = text == "Abstract"
        is_keywords = text.lower().startswith("keywords:")
        is_references_heading = text == "References"

        if is_references_heading:
            in_references = True
        elif is_heading and not is_references_heading:
            in_references = False

        if has_image:
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            para.style = doc.styles["Body Text"]
            apply_pagination_controls(para, widow=True, keep_with_next=True, keep_together=True)
            image_count += 1
            continue

        if not text:
            continue

        if is_title:
            para.style = doc.styles["Title"]
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(8)
            apply_pagination_controls(para, widow=True, keep_with_next=True, keep_together=True)
            continue

        if is_author:
            para.style = doc.styles["Subtitle"]
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(2)
            apply_pagination_controls(para, widow=True, keep_with_next=True, keep_together=True)
            continue

        if is_affiliation:
            para.style = doc.styles["Subtitle"]
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(8)
            for run in para.runs:
                run.font.size = Pt(config["affiliation_font_size"])
            apply_pagination_controls(para, widow=True, keep_with_next=False, keep_together=True)
            continue

        if is_abstract_label:
            para.style = doc.styles["Abstract Label"]
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            apply_pagination_controls(para, widow=True, keep_with_next=True, keep_together=True)
            continue

        if is_keywords:
            para.style = doc.styles["Keywords"]
            para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            if para.runs:
                para.runs[0].font.bold = True
            apply_pagination_controls(para, widow=True, keep_with_next=False, keep_together=False)
            continue

        if is_references_heading:
            para.style = doc.styles["References Heading"]
            para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            apply_pagination_controls(para, widow=True, keep_with_next=True, keep_together=True)
            continue

        if in_references:
            para.style = doc.styles["Reference Entry"]
            para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            apply_pagination_controls(para, widow=True, keep_with_next=False, keep_together=False)
            text_count += 1
            continue

        if is_caption:
            para.style = doc.styles["Table Caption"] if is_table_caption else doc.styles["Image Caption"]
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            emphasize_caption_label(para)
            set_spacing(
                para,
                before=config["table_space_before_caption"] if is_table_caption else 4,
                after=3 if is_table_caption else 6,
                line_spacing=1.0,
            )
            apply_pagination_controls(
                para,
                widow=True,
                keep_with_next=is_table_caption,
                keep_together=True,
            )
            continue

        if is_heading:
            apply_pagination_controls(para, widow=True, keep_with_next=True, keep_together=True)
            continue

        if style_name == "Compact":
            para.style = doc.styles["Compact"]
            para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            if previous_block_is_table(para):
                set_spacing(para, before=config["table_space_after"])
            apply_pagination_controls(para, widow=True, keep_with_next=False, keep_together=False)
            text_count += 1
            continue

        if style_name == "First Paragraph":
            para.style = doc.styles["First Paragraph"]
        else:
            para.style = doc.styles["Body Text"]

        para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        if has_display_math(para):
            set_spacing(
                para,
                before=config["equation_space_before"],
                after=config["equation_space_after"],
                line_spacing=1.0,
            )
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            apply_pagination_controls(para, widow=True, keep_with_next=False, keep_together=True)
            text_count += 1
            continue
        if previous_block_is_table(para):
            set_spacing(para, before=config["table_space_after"])
        apply_pagination_controls(para, widow=True, keep_with_next=False, keep_together=False)
        text_count += 1

    print(f"  Centered {image_count} images")
    print(f"  Formatted {text_count} text paragraphs")

    print("  Formatting tables...")
    for idx, table in enumerate(doc.tables):
        print(f"    Table {idx + 1}: {len(table.rows)} rows x {len(table.columns)} cols")
        set_table_width(table, table_width)
        set_table_layout(table)
        set_table_border(table, config)

        last_row_idx = len(table.rows) - 1

        if table.rows:
            for col_idx, cell in enumerate(table.rows[0].cells):
                format_table_cell(
                    cell,
                    config,
                    is_header=True,
                    alignment=WD_PARAGRAPH_ALIGNMENT.CENTER,
                    bottom_color=config["table_rule_mid"],
                    bottom_size=8,
                )

            for row_idx, row in enumerate(table.rows[1:], start=1):
                for col_idx, cell in enumerate(row.cells):
                    bottom_color = None if row_idx == last_row_idx else config["table_rule_light"]
                    format_table_cell(
                        cell,
                        config,
                        is_header=False,
                        alignment=WD_PARAGRAPH_ALIGNMENT.CENTER,
                        bottom_color=bottom_color,
                        bottom_size=4,
                    )

    print("  Formatting complete")


def main():
    parser = argparse.ArgumentParser(description="Apply academic paper formatting to DOCX")
    parser.add_argument("--input", "-i", default="paper_converted.docx", help="Input DOCX file")
    parser.add_argument("--output", "-o", default="paper_academic.docx", help="Output DOCX file")
    parser.add_argument("--table-width", type=int, default=100, help="Table width percentage")
    parser.add_argument(
        "--profile",
        default="camera-ready-generic",
        choices=sorted(PROFILES.keys()),
        help="Formatting profile",
    )
    args = parser.parse_args()

    if not os.path.exists(args.input):
        print(f"Error: Input file '{args.input}' not found")
        sys.exit(1)

    print(f"Loading: {args.input}")
    doc = Document(args.input)
    apply_academic_formatting(doc, profile=args.profile, table_width=args.table_width)
    doc.save(args.output)
    print(f"\nSaved: {args.output}")
    print("Done!")


if __name__ == "__main__":
    main()
